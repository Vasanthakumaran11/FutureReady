import io
import fitz  # PyMuPDF
import docx
from typing import Tuple

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB


class ExtractionError(Exception):
    def __init__(self, error_code: str, message: str):
        self.error_code = error_code
        self.message = message
        super().__init__(message)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text page-by-page from PDF bytes using PyMuPDF."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                pages_text.append(text.strip())
        doc.close()
        return "\n\n".join(pages_text)
    except Exception as e:
        raise ExtractionError("file_read_failed", f"Failed to read PDF document: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from paragraphs and tables in DOCX document order."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs_text = []

        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                paragraphs_text.append(p.text.strip())

        # Extract table text (e.g. skills grids)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    paragraphs_text.append(" | ".join(row_cells))

        return "\n\n".join(paragraphs_text)
    except Exception as e:
        raise ExtractionError("file_read_failed", f"Failed to read DOCX document: {str(e)}")


def process_resume_file(filename: str, file_bytes: bytes) -> Tuple[str, int]:
    """
    Validates file format, enforces 5MB limit, extracts raw text,
    and checks for scanned/unextractable text.
    """
    size_bytes = len(file_bytes)
    if size_bytes > MAX_FILE_SIZE_BYTES:
        raise ExtractionError("file_too_large", "Resume file exceeds maximum allowed limit of 5MB.")

    ext = filename.lower().split(".")[-1]
    if ext not in ["pdf", "docx"]:
        raise ExtractionError(
            "unsupported_file_type", "Only .pdf and .docx resume files are supported."
        )

    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    else:
        raw_text = extract_text_from_docx(file_bytes)

    cleaned_text = raw_text.strip()

    # Check for scanned image PDF (<50 chars)
    if len(cleaned_text) < 50:
        raise ExtractionError(
            "no_extractable_text",
            "This looks like a scanned or image-based file. Please upload a text-based PDF/DOCX, or use the guided builder instead.",
        )

    return cleaned_text, size_bytes
